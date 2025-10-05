"""
File conversion utilities
"""
import os
import json
import zipfile
import subprocess
from pathlib import Path
from PIL import Image
from pdf2docx import Converter as PDFConverter
from docx import Document
from docx.shared import Inches
import img2pdf
from pydub import AudioSegment
import pandas as pd
import markdown
from io import BytesIO
import tempfile

class FileConverter:
    """Main file converter class"""
    
    # Define conversion mappings
    CONVERSION_MAP = {
        'pdf': ['docx', 'txt', 'html', 'jpg', 'png'],
        'docx': ['pdf', 'txt', 'html'],
        'txt': ['pdf', 'docx'],
        'pptx': ['pdf', 'jpg', 'png'],
        'xlsx': ['pdf', 'txt', 'json', 'csv'],
        'csv': ['pdf', 'txt', 'json', 'xlsx'],
        'jpg': ['png', 'webp', 'pdf', 'bmp', 'tiff'],
        'jpeg': ['png', 'webp', 'pdf', 'bmp', 'tiff'],
        'png': ['jpg', 'webp', 'pdf', 'bmp', 'tiff'],
        'webp': ['jpg', 'png', 'pdf'],
        'bmp': ['jpg', 'png', 'pdf'],
        'tiff': ['jpg', 'png', 'pdf'],
        'svg': ['png', 'jpg'],
        'mp3': ['wav', 'ogg', 'flac', 'aac'],
        'wav': ['mp3', 'ogg', 'flac', 'aac'],
        'ogg': ['mp3', 'wav', 'flac', 'aac'],
        'flac': ['mp3', 'wav', 'ogg', 'aac'],
        'aac': ['mp3', 'wav', 'ogg', 'flac'],
        'mp4': ['avi', 'mov', 'mkv', 'gif', 'mp3'],
        'avi': ['mp4', 'mov', 'mkv', 'mp3'],
        'mov': ['mp4', 'avi', 'mkv', 'mp3'],
        'mkv': ['mp4', 'avi', 'mov', 'mp3'],
        'gif': ['mp4'],
        'zip': ['extract'],
        'rar': ['extract'],
        '7z': ['extract'],
        'md': ['pdf', 'html'],
        'json': ['csv'],
    }

    @staticmethod
    def get_valid_conversions(file_extension):
        """Get valid conversion formats for a file type"""
        ext = file_extension.lower().lstrip('.')
        return FileConverter.CONVERSION_MAP.get(ext, [])

    @staticmethod
    def detect_file_type(filename):
        """Detect file type from extension"""
        return Path(filename).suffix.lower().lstrip('.')

    @staticmethod
    def sanitize_filename(filename):
        """Sanitize filename for security"""
        import re
        # Remove any path components
        filename = os.path.basename(filename)
        # Remove any non-alphanumeric characters except dots, hyphens, and underscores
        filename = re.sub(r'[^\w\.-]', '_', filename)
        return filename

    @staticmethod
    def convert(input_path, output_path, target_format):
        """Main conversion dispatcher"""
        source_ext = FileConverter.detect_file_type(input_path)
        target_ext = target_format.lower()

        # Document conversions
        if source_ext == 'pdf' and target_ext == 'docx':
            return FileConverter._pdf_to_docx(input_path, output_path)
        elif source_ext == 'pdf' and target_ext == 'txt':
            return FileConverter._pdf_to_txt(input_path, output_path)
        elif source_ext == 'pdf' and target_ext == 'html':
            return FileConverter._pdf_to_html(input_path, output_path)
        elif source_ext == 'pdf' and target_ext in ['jpg', 'png']:
            return FileConverter._pdf_to_images(input_path, output_path, target_ext)
        elif source_ext == 'docx' and target_ext == 'pdf':
            return FileConverter._docx_to_pdf(input_path, output_path)
        elif source_ext == 'docx' and target_ext == 'txt':
            return FileConverter._docx_to_txt(input_path, output_path)
        elif source_ext == 'docx' and target_ext == 'html':
            return FileConverter._docx_to_html(input_path, output_path)
        elif source_ext == 'txt' and target_ext == 'pdf':
            return FileConverter._txt_to_pdf(input_path, output_path)
        elif source_ext == 'txt' and target_ext == 'docx':
            return FileConverter._txt_to_docx(input_path, output_path)
        
        # Excel/CSV conversions
        elif source_ext in ['xlsx', 'csv'] and target_ext in ['pdf', 'txt', 'json', 'csv', 'xlsx']:
            return FileConverter._spreadsheet_convert(input_path, output_path, source_ext, target_ext)
        
        # Image conversions
        elif source_ext in ['jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff']:
            if target_ext == 'pdf':
                return FileConverter._image_to_pdf(input_path, output_path)
            else:
                return FileConverter._image_convert(input_path, output_path, target_ext)
        
        # SVG conversions
        elif source_ext == 'svg' and target_ext in ['png', 'jpg']:
            return FileConverter._svg_to_image(input_path, output_path, target_ext)
        
        # Audio conversions
        elif source_ext in ['mp3', 'wav', 'ogg', 'flac', 'aac']:
            return FileConverter._audio_convert(input_path, output_path, target_ext)
        
        # Video conversions
        elif source_ext in ['mp4', 'avi', 'mov', 'mkv']:
            if target_ext == 'mp3':
                return FileConverter._video_to_audio(input_path, output_path)
            elif target_ext == 'gif':
                return FileConverter._video_to_gif(input_path, output_path)
            else:
                return FileConverter._video_convert(input_path, output_path, target_ext)
        elif source_ext == 'gif' and target_ext == 'mp4':
            return FileConverter._gif_to_video(input_path, output_path)
        
        # Markdown conversions
        elif source_ext == 'md' and target_ext == 'html':
            return FileConverter._md_to_html(input_path, output_path)
        elif source_ext == 'md' and target_ext == 'pdf':
            return FileConverter._md_to_pdf(input_path, output_path)
        
        # JSON conversions
        elif source_ext == 'json' and target_ext == 'csv':
            return FileConverter._json_to_csv(input_path, output_path)
        
        else:
            raise ValueError(f"Conversion from {source_ext} to {target_ext} not supported")

    # Document conversion methods
    @staticmethod
    def _pdf_to_docx(input_path, output_path):
        cv = PDFConverter(input_path)
        cv.convert(output_path)
        cv.close()
        return True

    @staticmethod
    def _pdf_to_txt(input_path, output_path):
        import PyPDF2
        with open(input_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = []
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        with open(output_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write('\n\n'.join(text))
        return True

    @staticmethod
    def _pdf_to_html(input_path, output_path):
        import PyPDF2
        with open(input_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = []
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        html_content = '<html><body>\n'
        for i, page_text in enumerate(text):
            html_content += f'<div class="page" id="page-{i+1}">\n'
            html_content += f'<p>{page_text.replace(chr(10), "<br>")}</p>\n'
            html_content += '</div>\n'
        html_content += '</body></html>'
        with open(output_path, 'w', encoding='utf-8') as html_file:
            html_file.write(html_content)
        return True

    @staticmethod
    def _pdf_to_images(input_path, output_path, format):
        from pdf2image import convert_from_path
        images = convert_from_path(input_path)
        if len(images) == 1:
            images[0].save(output_path, format.upper())
        else:
            # Save first page only
            images[0].save(output_path, format.upper())
        return True

    @staticmethod
    def _docx_to_pdf(input_path, output_path):
        # Use LibreOffice for conversion
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_path), input_path
        ], check=True)
        # Move the output file to the correct location
        temp_output = os.path.join(os.path.dirname(output_path), 
                                   Path(input_path).stem + '.pdf')
        if temp_output != output_path:
            os.rename(temp_output, output_path)
        return True

    @staticmethod
    def _docx_to_txt(input_path, output_path):
        doc = Document(input_path)
        with open(output_path, 'w', encoding='utf-8') as txt_file:
            for para in doc.paragraphs:
                txt_file.write(para.text + '\n')
        return True

    @staticmethod
    def _docx_to_html(input_path, output_path):
        doc = Document(input_path)
        html_content = '<html><body>\n'
        for para in doc.paragraphs:
            html_content += f'<p>{para.text}</p>\n'
        html_content += '</body></html>'
        with open(output_path, 'w', encoding='utf-8') as html_file:
            html_file.write(html_content)
        return True

    @staticmethod
    def _txt_to_pdf(input_path, output_path):
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(output_path, pagesize=letter)
        with open(input_path, 'r', encoding='utf-8') as txt_file:
            text = txt_file.read()
            y = 750
            for line in text.split('\n'):
                if y < 50:
                    c.showPage()
                    y = 750
                c.drawString(50, y, line[:100])
                y -= 15
        c.save()
        return True

    @staticmethod
    def _txt_to_docx(input_path, output_path):
        doc = Document()
        with open(input_path, 'r', encoding='utf-8') as txt_file:
            for line in txt_file:
                doc.add_paragraph(line.rstrip())
        doc.save(output_path)
        return True

    # Spreadsheet conversions
    @staticmethod
    def _spreadsheet_convert(input_path, output_path, source_ext, target_ext):
        if source_ext == 'csv':
            df = pd.read_csv(input_path)
        else:
            df = pd.read_excel(input_path)
        
        if target_ext == 'csv':
            df.to_csv(output_path, index=False)
        elif target_ext == 'xlsx':
            df.to_excel(output_path, index=False)
        elif target_ext == 'json':
            df.to_json(output_path, orient='records', indent=2)
        elif target_ext == 'txt':
            with open(output_path, 'w') as f:
                f.write(df.to_string())
        elif target_ext == 'pdf':
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(output_path, pagesize=landscape(letter))
            data = [df.columns.tolist()] + df.values.tolist()
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            doc.build([table])
        return True

    # Image conversions
    @staticmethod
    def _image_convert(input_path, output_path, target_format):
        img = Image.open(input_path)
        if img.mode == 'RGBA' and target_format.lower() in ['jpg', 'jpeg']:
            img = img.convert('RGB')
        img.save(output_path, format=target_format.upper())
        return True

    @staticmethod
    def _image_to_pdf(input_path, output_path):
        with open(output_path, 'wb') as f:
            f.write(img2pdf.convert(input_path))
        return True

    @staticmethod
    def _svg_to_image(input_path, output_path, target_format):
        import cairosvg
        if target_format == 'png':
            cairosvg.svg2png(url=input_path, write_to=output_path)
        elif target_format == 'jpg':
            # Convert to PNG first, then to JPG
            temp_png = output_path + '.temp.png'
            cairosvg.svg2png(url=input_path, write_to=temp_png)
            img = Image.open(temp_png)
            img = img.convert('RGB')
            img.save(output_path, 'JPEG')
            os.remove(temp_png)
        return True

    # Audio conversions
    @staticmethod
    def _audio_convert(input_path, output_path, target_format):
        audio = AudioSegment.from_file(input_path)
        audio.export(output_path, format=target_format)
        return True

    # Video conversions
    @staticmethod
    def _video_convert(input_path, output_path, target_format):
        cmd = ['ffmpeg', '-i', input_path, '-y', output_path]
        subprocess.run(cmd, check=True, capture_output=True)
        return True

    @staticmethod
    def _video_to_audio(input_path, output_path):
        cmd = ['ffmpeg', '-i', input_path, '-q:a', '0', '-map', 'a', '-y', output_path]
        subprocess.run(cmd, check=True, capture_output=True)
        return True

    @staticmethod
    def _video_to_gif(input_path, output_path):
        cmd = ['ffmpeg', '-i', input_path, '-vf', 'fps=10,scale=480:-1:flags=lanczos', 
               '-y', output_path]
        subprocess.run(cmd, check=True, capture_output=True)
        return True

    @staticmethod
    def _gif_to_video(input_path, output_path):
        cmd = ['ffmpeg', '-i', input_path, '-movflags', 'faststart', '-pix_fmt', 'yuv420p',
               '-vf', 'scale=trunc(iw/2)*2:trunc(ih/2)*2', '-y', output_path]
        subprocess.run(cmd, check=True, capture_output=True)
        return True

    # Markdown conversions
    @staticmethod
    def _md_to_html(input_path, output_path):
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        html_content = markdown.markdown(md_content)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f'<html><body>{html_content}</body></html>')
        return True

    @staticmethod
    def _md_to_pdf(input_path, output_path):
        # Convert to HTML first, then to PDF
        temp_html = output_path + '.temp.html'
        FileConverter._md_to_html(input_path, temp_html)
        
        import pdfkit
        pdfkit.from_file(temp_html, output_path)
        os.remove(temp_html)
        return True

    # JSON conversions
    @staticmethod
    def _json_to_csv(input_path, output_path):
        with open(input_path, 'r') as f:
            data = json.load(f)
        df = pd.DataFrame(data)
        df.to_csv(output_path, index=False)
        return True
