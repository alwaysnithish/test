from django.shortcuts import render
from django.shortcuts import render, redirect
from django.http import JsonResponse, HttpResponse, FileResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings
import json
import os
import io
import zipfile
from datetime import datetime
import tempfile

# PDF processing libraries
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
from PIL import Image
import fitz  # PyMuPDF for advanced operations
from pdf2image import convert_from_path, convert_from_bytes
import pdfplumber
import pytesseract
from fpdf import FPDF

# Utility functions
def save_uploaded_file(uploaded_file):
    """Save uploaded file temporarily and return path"""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, uploaded_file.name)
    
    with open(file_path, 'wb+') as destination:
        for chunk in uploaded_file.chunks():
            destination.write(chunk)
    
    return file_path

def get_pdf_info(file_path):
    """Extract basic PDF information"""
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            info = {
                'pages': len(reader.pages),
                'title': reader.metadata.get('/Title', 'Unknown') if reader.metadata else 'Unknown',
                'author': reader.metadata.get('/Author', 'Unknown') if reader.metadata else 'Unknown',
                'subject': reader.metadata.get('/Subject', 'Unknown') if reader.metadata else 'Unknown',
                'creator': reader.metadata.get('/Creator', 'Unknown') if reader.metadata else 'Unknown',
                'producer': reader.metadata.get('/Producer', 'Unknown') if reader.metadata else 'Unknown',
                'creation_date': reader.metadata.get('/CreationDate', 'Unknown') if reader.metadata else 'Unknown',
            }
            return info
    except Exception as e:
        return {'error': str(e)}

# Main views
def pdf_tools_home(request):
    """Home page for PDF tools"""
    context = {
        'title': 'Universal PDF Tools',
        'tools': [
            'Extract Text', 'Extract Images', 'Split PDF', 'Merge PDFs',
            'Compress PDF', 'Convert to Images', 'Add Watermark', 'Remove Password',
            'Add Password', 'Rotate Pages', 'Delete Pages', 'Rearrange Pages',
            'PDF to Word', 'Images to PDF', 'HTML to PDF', 'Extract Pages'
        ]
    }
    return render(request, 'pdftools.html', context)

@csrf_exempt
@require_http_methods(["POST"])
def upload_pdf(request):
    """Handle PDF file upload and return basic info"""
    try:
        if 'pdf_file' not in request.FILES:
            return JsonResponse({'error': 'No PDF file uploaded'}, status=400)
        
        pdf_file = request.FILES['pdf_file']
        
        if not pdf_file.name.lower().endswith('.pdf'):
            return JsonResponse({'error': 'Please upload a valid PDF file'}, status=400)
        
        # Save file temporarily
        file_path = save_uploaded_file(pdf_file)
        
        # Get PDF info
        pdf_info = get_pdf_info(file_path)
        
        # Store file path in session for subsequent operations
        request.session['current_pdf'] = file_path
        request.session['original_filename'] = pdf_file.name
        
        return JsonResponse({
            'success': True,
            'filename': pdf_file.name,
            'info': pdf_info,
            'operations_available': True
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def extract_text(request):
    """Extract text from PDF"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        extracted_text = []
        
        # Method 1: Try PyPDF2 first
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    extracted_text.append({
                        'page': page_num,
                        'text': text,
                        'method': 'PyPDF2'
                    })
        except:
            # Method 2: Use pdfplumber as fallback
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text() or ""
                        extracted_text.append({
                            'page': page_num,
                            'text': text,
                            'method': 'pdfplumber'
                        })
            except:
                # Method 3: OCR as last resort
                try:
                    doc = fitz.open(file_path)
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        text = pytesseract.image_to_string(img)
                        extracted_text.append({
                            'page': page_num + 1,
                            'text': text,
                            'method': 'OCR'
                        })
                    doc.close()
                except Exception as e:
                    return JsonResponse({'error': f'Text extraction failed: {str(e)}'}, status=500)
        
        # Combine all text
        full_text = "\n\n".join([f"Page {item['page']}:\n{item['text']}" for item in extracted_text])
        
        return JsonResponse({
            'success': True,
            'text': full_text,
            'pages': extracted_text,
            'total_pages': len(extracted_text)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def extract_images(request):
    """Extract images from PDF"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        images = []
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    
                    # Save image temporarily
                    temp_dir = tempfile.mkdtemp()
                    img_filename = f"page_{page_num + 1}_img_{img_index + 1}.png"
                    img_path = os.path.join(temp_dir, img_filename)
                    
                    with open(img_path, 'wb') as img_file:
                        img_file.write(img_data)
                    
                    images.append({
                        'page': page_num + 1,
                        'image_index': img_index + 1,
                        'filename': img_filename,
                        'path': img_path,
                        'size': len(img_data)
                    })
                
                pix = None
        
        doc.close()
        
        return JsonResponse({
            'success': True,
            'images': images,
            'total_images': len(images)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def split_pdf(request):
    """Split PDF into individual pages or ranges"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body)
        split_type = data.get('split_type', 'pages')  # 'pages' or 'ranges'
        ranges = data.get('ranges', [])  # List of ranges like [{'start': 1, 'end': 5}]
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            total_pages = len(reader.pages)
            
            split_files = []
            
            if split_type == 'pages':
                # Split into individual pages
                for page_num in range(total_pages):
                    writer = PdfWriter()
                    writer.add_page(reader.pages[page_num])
                    
                    # Save split file
                    temp_dir = tempfile.mkdtemp()
                    split_filename = f"page_{page_num + 1}.pdf"
                    split_path = os.path.join(temp_dir, split_filename)
                    
                    with open(split_path, 'wb') as output_file:
                        writer.write(output_file)
                    
                    split_files.append({
                        'filename': split_filename,
                        'path': split_path,
                        'pages': f"{page_num + 1}"
                    })
            
            elif split_type == 'ranges' and ranges:
                # Split by ranges
                for i, page_range in enumerate(ranges):
                    start = max(1, page_range.get('start', 1)) - 1  # Convert to 0-based
                    end = min(total_pages, page_range.get('end', total_pages))
                    
                    writer = PdfWriter()
                    for page_num in range(start, end):
                        writer.add_page(reader.pages[page_num])
                    
                    # Save split file
                    temp_dir = tempfile.mkdtemp()
                    split_filename = f"pages_{start + 1}_to_{end}.pdf"
                    split_path = os.path.join(temp_dir, split_filename)
                    
                    with open(split_path, 'wb') as output_file:
                        writer.write(output_file)
                    
                    split_files.append({
                        'filename': split_filename,
                        'path': split_path,
                        'pages': f"{start + 1}-{end}"
                    })
        
        return JsonResponse({
            'success': True,
            'split_files': split_files,
            'total_files': len(split_files)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def merge_pdfs(request):
    """Merge multiple PDF files"""
    try:
        if 'pdf_files' not in request.FILES:
            return JsonResponse({'error': 'No PDF files uploaded'}, status=400)
        
        pdf_files = request.FILES.getlist('pdf_files')
        
        if len(pdf_files) < 2:
            return JsonResponse({'error': 'At least 2 PDF files required for merging'}, status=400)
        
        writer = PdfWriter()
        
        for pdf_file in pdf_files:
            if not pdf_file.name.lower().endswith('.pdf'):
                continue
                
            file_path = save_uploaded_file(pdf_file)
            
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    writer.add_page(page)
        
        # Save merged file
        temp_dir = tempfile.mkdtemp()
        merged_filename = f"merged_pdf_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        merged_path = os.path.join(temp_dir, merged_filename)
        
        with open(merged_path, 'wb') as output_file:
            writer.write(output_file)
        
        return JsonResponse({
            'success': True,
            'merged_file': {
                'filename': merged_filename,
                'path': merged_path
            },
            'total_input_files': len(pdf_files)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def compress_pdf(request):
    """Compress PDF file"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        # Get original file size
        original_size = os.path.getsize(file_path)
        
        # Open and rewrite PDF (basic compression)
        doc = fitz.open(file_path)
        
        temp_dir = tempfile.mkdtemp()
        compressed_filename = f"compressed_{request.session.get('original_filename', 'file.pdf')}"
        compressed_path = os.path.join(temp_dir, compressed_filename)
        
        # Save with compression
        doc.save(compressed_path, garbage=4, deflate=True, clean=True)
        doc.close()
        
        # Get compressed file size
        compressed_size = os.path.getsize(compressed_path)
        compression_ratio = (original_size - compressed_size) / original_size * 100
        
        return JsonResponse({
            'success': True,
            'compressed_file': {
                'filename': compressed_filename,
                'path': compressed_path
            },
            'original_size': original_size,
            'compressed_size': compressed_size,
            'compression_ratio': round(compression_ratio, 2)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def convert_to_images(request):
    """Convert PDF pages to images"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body) if request.body else {}
        image_format = data.get('format', 'PNG').upper()
        dpi = data.get('dpi', 200)
        
        # Convert PDF to images
        images = convert_from_path(file_path, dpi=dpi)
        
        image_files = []
        temp_dir = tempfile.mkdtemp()
        
        for i, image in enumerate(images):
            image_filename = f"page_{i + 1}.{image_format.lower()}"
            image_path = os.path.join(temp_dir, image_filename)
            
            image.save(image_path, image_format)
            
            image_files.append({
                'page': i + 1,
                'filename': image_filename,
                'path': image_path,
                'format': image_format,
                'size': os.path.getsize(image_path)
            })
        
        return JsonResponse({
            'success': True,
            'images': image_files,
            'total_images': len(image_files),
            'format': image_format,
            'dpi': dpi
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def add_watermark(request):
    """Add watermark to PDF"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body)
        watermark_text = data.get('text', 'WATERMARK')
        opacity = data.get('opacity', 0.3)
        position = data.get('position', 'center')  # center, top-left, top-right, bottom-left, bottom-right
        
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # Get page dimensions
            page_rect = page.rect
            
            # Set position
            if position == 'center':
                x = page_rect.width / 2
                y = page_rect.height / 2
            elif position == 'top-left':
                x, y = 50, 50
            elif position == 'top-right':
                x, y = page_rect.width - 150, 50
            elif position == 'bottom-left':
                x, y = 50, page_rect.height - 50
            elif position == 'bottom-right':
                x, y = page_rect.width - 150, page_rect.height - 50
            else:
                x = page_rect.width / 2
                y = page_rect.height / 2
            
            # Add text watermark
            page.insert_text(
                (x, y),
                watermark_text,
                fontsize=36,
                color=(0.5, 0.5, 0.5),
                rotate=45,
                overlay=True
            )
        
        # Save watermarked PDF
        temp_dir = tempfile.mkdtemp()
        watermarked_filename = f"watermarked_{request.session.get('original_filename', 'file.pdf')}"
        watermarked_path = os.path.join(temp_dir, watermarked_filename)
        
        doc.save(watermarked_path)
        doc.close()
        
        return JsonResponse({
            'success': True,
            'watermarked_file': {
                'filename': watermarked_filename,
                'path': watermarked_path
            },
            'watermark_text': watermark_text,
            'position': position
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def add_password(request):
    """Add password protection to PDF"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body)
        password = data.get('password', '')
        
        if not password:
            return JsonResponse({'error': 'Password is required'}, status=400)
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            writer = PdfWriter()
            
            # Copy all pages
            for page in reader.pages:
                writer.add_page(page)
            
            # Add password
            writer.encrypt(password)
            
            # Save password-protected PDF
            temp_dir = tempfile.mkdtemp()
            protected_filename = f"password_protected_{request.session.get('original_filename', 'file.pdf')}"
            protected_path = os.path.join(temp_dir, protected_filename)
            
            with open(protected_path, 'wb') as output_file:
                writer.write(output_file)
        
        return JsonResponse({
            'success': True,
            'protected_file': {
                'filename': protected_filename,
                'path': protected_path
            }
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def remove_password(request):
    """Remove password from PDF"""
    try:
        if 'pdf_file' not in request.FILES:
            return JsonResponse({'error': 'No PDF file uploaded'}, status=400)
        
        pdf_file = request.FILES['pdf_file']
        data = json.loads(request.POST.get('data', '{}'))
        password = data.get('password', '')
        
        if not password:
            return JsonResponse({'error': 'Password is required'}, status=400)
        
        file_path = save_uploaded_file(pdf_file)
        
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                
                if reader.is_encrypted:
                    reader.decrypt(password)
                
                writer = PdfWriter()
                
                # Copy all pages
                for page in reader.pages:
                    writer.add_page(page)
                
                # Save unprotected PDF
                temp_dir = tempfile.mkdtemp()
                unprotected_filename = f"unprotected_{pdf_file.name}"
                unprotected_path = os.path.join(temp_dir, unprotected_filename)
                
                with open(unprotected_path, 'wb') as output_file:
                    writer.write(output_file)
            
            return JsonResponse({
                'success': True,
                'unprotected_file': {
                    'filename': unprotected_filename,
                    'path': unprotected_path
                }
            })
            
        except Exception as decrypt_error:
            return JsonResponse({'error': 'Invalid password or cannot decrypt PDF'}, status=400)
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def rotate_pages(request):
    """Rotate PDF pages"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body)
        rotation = data.get('rotation', 90)  # 90, 180, 270, -90
        pages = data.get('pages', 'all')  # 'all' or list of page numbers
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            writer = PdfWriter()
            
            total_pages = len(reader.pages)
            
            if pages == 'all':
                pages_to_rotate = list(range(total_pages))
            else:
                pages_to_rotate = [p - 1 for p in pages if 1 <= p <= total_pages]  # Convert to 0-based
            
            for page_num in range(total_pages):
                page = reader.pages[page_num]
                
                if page_num in pages_to_rotate:
                    page.rotate(rotation)
                
                writer.add_page(page)
            
            # Save rotated PDF
            temp_dir = tempfile.mkdtemp()
            rotated_filename = f"rotated_{request.session.get('original_filename', 'file.pdf')}"
            rotated_path = os.path.join(temp_dir, rotated_filename)
            
            with open(rotated_path, 'wb') as output_file:
                writer.write(output_file)
        
        return JsonResponse({
            'success': True,
            'rotated_file': {
                'filename': rotated_filename,
                'path': rotated_path
            },
            'rotation': rotation,
            'pages_rotated': len(pages_to_rotate)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def delete_pages(request):
    """Delete specific pages from PDF"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body)
        pages_to_delete = data.get('pages', [])  # List of page numbers to delete
        
        if not pages_to_delete:
            return JsonResponse({'error': 'No pages specified for deletion'}, status=400)
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            writer = PdfWriter()
            
            total_pages = len(reader.pages)
            pages_to_delete_zero_based = [p - 1 for p in pages_to_delete if 1 <= p <= total_pages]
            
            for page_num in range(total_pages):
                if page_num not in pages_to_delete_zero_based:
                    writer.add_page(reader.pages[page_num])
            
            if len(writer.pages) == 0:
                return JsonResponse({'error': 'Cannot delete all pages'}, status=400)
            
            # Save PDF with deleted pages
            temp_dir = tempfile.mkdtemp()
            modified_filename = f"deleted_pages_{request.session.get('original_filename', 'file.pdf')}"
            modified_path = os.path.join(temp_dir, modified_filename)
            
            with open(modified_path, 'wb') as output_file:
                writer.write(output_file)
        
        return JsonResponse({
            'success': True,
            'modified_file': {
                'filename': modified_filename,
                'path': modified_path
            },
            'original_pages': total_pages,
            'remaining_pages': len(writer.pages),
            'deleted_pages': len(pages_to_delete_zero_based)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def images_to_pdf(request):
    """Convert images to PDF"""
    try:
        if 'image_files' not in request.FILES:
            return JsonResponse({'error': 'No image files uploaded'}, status=400)
        
        image_files = request.FILES.getlist('image_files')
        
        if not image_files:
            return JsonResponse({'error': 'No image files provided'}, status=400)
        
        pdf = FPDF()
        
        for image_file in image_files:
            # Save image temporarily
            temp_dir = tempfile.mkdtemp()
            image_path = os.path.join(temp_dir, image_file.name)
            
            with open(image_path, 'wb+') as destination:
                for chunk in image_file.chunks():
                    destination.write(chunk)
            
            # Add image to PDF
            try:
                img = Image.open(image_path)
                img_width, img_height = img.size
                
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                    rgb_path = image_path.replace(image_file.name, f"rgb_{image_file.name}")
                    img.save(rgb_path)
                    image_path = rgb_path
                
                pdf.add_page()
                
                # Calculate dimensions to fit page
                page_width = 210  # A4 width in mm
                page_height = 297  # A4 height in mm
                
                aspect_ratio = img_width / img_height
                
                if aspect_ratio > page_width / page_height:
                    # Image is wider
                    width = page_width - 20  # 10mm margin on each side
                    height = width / aspect_ratio
                else:
                    # Image is taller
                    height = page_height - 20  # 10mm margin on top and bottom
                    width = height * aspect_ratio
                
                # Center the image
                x = (page_width - width) / 2
                y = (page_height - height) / 2
                
                pdf.image(image_path, x, y, width, height)
                
            except Exception as img_error:
                continue  # Skip problematic images
        
        if pdf.page_count() == 0:
            return JsonResponse({'error': 'No valid images could be processed'}, status=400)
        
        # Save PDF
        temp_dir = tempfile.mkdtemp()
        pdf_filename = f"images_to_pdf_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        pdf.output(pdf_path, 'F')
        
        return JsonResponse({
            'success': True,
            'pdf_file': {
                'filename': pdf_filename,
                'path': pdf_path
            },
            'total_images': len(image_files),
            'processed_images': pdf.page_count()
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def download_file(request, file_type):
    """Download processed file"""
    try:
        file_path = request.GET.get('path', '')
        filename = request.GET.get('filename', 'download.pdf')
        
        if not file_path or not os.path.exists(file_path):
            return HttpResponse('File not found', status=404)
        
        response = FileResponse(
            open(file_path, 'rb'),
            as_attachment=True,
            filename=filename
        )
        
        return response
        
    except Exception as e:
        return HttpResponse(f'Error downloading file: {str(e)}', status=500)

def download_zip(request):
    """Download multiple files as ZIP"""
    try:
        file_paths = request.GET.getlist('paths')
        filenames = request.GET.getlist('filenames')
        
        if not file_paths:
            return HttpResponse('No files specified', status=400)
        
        # Create ZIP file
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for i, file_path in enumerate(file_paths):
                if os.path.exists(file_path):
                    filename = filenames[i] if i < len(filenames) else f"file_{i+1}.pdf"
                    zip_file.write(file_path, filename)
        
        zip_buffer.seek(0)
        
        response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="pdf_tools_files_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip"'
        
        return response
        
    except Exception as e:
        return HttpResponse(f'Error creating ZIP file: {str(e)}', status=500)

@csrf_exempt
@require_http_methods(["POST"])
def extract_pages(request):
    """Extract specific pages from PDF"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body)
        pages_to_extract = data.get('pages', [])  # List of page numbers to extract
        
        if not pages_to_extract:
            return JsonResponse({'error': 'No pages specified for extraction'}, status=400)
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            writer = PdfWriter()
            
            total_pages = len(reader.pages)
            valid_pages = [p - 1 for p in pages_to_extract if 1 <= p <= total_pages]  # Convert to 0-based
            
            for page_num in valid_pages:
                writer.add_page(reader.pages[page_num])
            
            if len(writer.pages) == 0:
                return JsonResponse({'error': 'No valid pages to extract'}, status=400)
            
            # Save extracted pages
            temp_dir = tempfile.mkdtemp()
            extracted_filename = f"extracted_pages_{request.session.get('original_filename', 'file.pdf')}"
            extracted_path = os.path.join(temp_dir, extracted_filename)
            
            with open(extracted_path, 'wb') as output_file:
                writer.write(output_file)
        
        return JsonResponse({
            'success': True,
            'extracted_file': {
                'filename': extracted_filename,
                'path': extracted_path
            },
            'original_pages': total_pages,
            'extracted_pages': len(valid_pages)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def rearrange_pages(request):
    """Rearrange pages in PDF"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body)
        page_order = data.get('page_order', [])  # List of page numbers in desired order
        
        if not page_order:
            return JsonResponse({'error': 'No page order specified'}, status=400)
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            writer = PdfWriter()
            
            total_pages = len(reader.pages)
            
            # Validate page order
            valid_pages = [p - 1 for p in page_order if 1 <= p <= total_pages]  # Convert to 0-based
            
            # Add pages in the specified order
            for page_num in valid_pages:
                writer.add_page(reader.pages[page_num])
            
            if len(writer.pages) == 0:
                return JsonResponse({'error': 'No valid pages in the specified order'}, status=400)
            
            # Save rearranged PDF
            temp_dir = tempfile.mkdtemp()
            rearranged_filename = f"rearranged_{request.session.get('original_filename', 'file.pdf')}"
            rearranged_path = os.path.join(temp_dir, rearranged_filename)
            
            with open(rearranged_path, 'wb') as output_file:
                writer.write(output_file)
        
        return JsonResponse({
            'success': True,
            'rearranged_file': {
                'filename': rearranged_filename,
                'path': rearranged_path
            },
            'original_pages': total_pages,
            'new_order': page_order
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def html_to_pdf(request):
    """Convert HTML content to PDF"""
    try:
        data = json.loads(request.body)
        html_content = data.get('html_content', '')
        
        if not html_content:
            return JsonResponse({'error': 'No HTML content provided'}, status=400)
        
        # Using weasyprint or pdfkit for HTML to PDF conversion
        # This is a simplified version - you might want to use libraries like:
        # - weasyprint
        # - pdfkit (requires wkhtmltopdf)
        # - reportlab with HTML parsing
        
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        import html2text
        
        # Convert HTML to plain text (basic conversion)
        h = html2text.HTML2Text()
        h.ignore_links = True
        text_content = h.handle(html_content)
        
        # Create PDF
        temp_dir = tempfile.mkdtemp()
        pdf_filename = f"html_to_pdf_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add content as paragraphs
        for line in text_content.split('\n'):
            if line.strip():
                para = Paragraph(line, styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
        
        return JsonResponse({
            'success': True,
            'pdf_file': {
                'filename': pdf_filename,
                'path': pdf_path
            }
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def pdf_to_word(request):
    """Convert PDF to Word document (basic text extraction)"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        # Extract text from PDF
        extracted_text = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    extracted_text.append(text)
        except:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    text = page.extract_text()
                    extracted_text.append(text)
        
        # Create Word document using python-docx
        from docx import Document
        
        doc = Document()
        doc.add_heading('Converted from PDF', 0)
        
        for i, page_text in enumerate(extracted_text, 1):
            if page_text.strip():
                doc.add_heading(f'Page {i}', level=1)
                doc.add_paragraph(page_text)
        
        # Save Word document
        temp_dir = tempfile.mkdtemp()
        word_filename = f"pdf_to_word_{request.session.get('original_filename', 'file.pdf').replace('.pdf', '.docx')}"
        word_path = os.path.join(temp_dir, word_filename)
        
        doc.save(word_path)
        
        return JsonResponse({
            'success': True,
            'word_file': {
                'filename': word_filename,
                'path': word_path
            },
            'total_pages': len(extracted_text)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def batch_process(request):
    """Process multiple PDFs with the same operation"""
    try:
        if 'pdf_files' not in request.FILES:
            return JsonResponse({'error': 'No PDF files uploaded'}, status=400)
        
        pdf_files = request.FILES.getlist('pdf_files')
        data = json.loads(request.POST.get('operation_data', '{}'))
        operation = data.get('operation', '')
        
        if not operation:
            return JsonResponse({'error': 'No operation specified'}, status=400)
        
        results = []
        
        for pdf_file in pdf_files:
            if not pdf_file.name.lower().endswith('.pdf'):
                continue
            
            try:
                file_path = save_uploaded_file(pdf_file)
                
                # Store in session temporarily for operation
                request.session['current_pdf'] = file_path
                request.session['original_filename'] = pdf_file.name
                
                # Perform operation based on type
                if operation == 'extract_text':
                    # Extract text
                    with pdfplumber.open(file_path) as pdf:
                        text = ""
                        for page in pdf.pages:
                            text += page.extract_text() or ""
                    
                    results.append({
                        'filename': pdf_file.name,
                        'operation': 'extract_text',
                        'result': 'success',
                        'text_length': len(text)
                    })
                
                elif operation == 'compress':
                    # Compress PDF
                    doc = fitz.open(file_path)
                    temp_dir = tempfile.mkdtemp()
                    compressed_path = os.path.join(temp_dir, f"compressed_{pdf_file.name}")
                    
                    doc.save(compressed_path, garbage=4, deflate=True, clean=True)
                    doc.close()
                    
                    original_size = os.path.getsize(file_path)
                    compressed_size = os.path.getsize(compressed_path)
                    
                    results.append({
                        'filename': pdf_file.name,
                        'operation': 'compress',
                        'result': 'success',
                        'original_size': original_size,
                        'compressed_size': compressed_size,
                        'compression_ratio': round((original_size - compressed_size) / original_size * 100, 2),
                        'output_path': compressed_path
                    })
                
                elif operation == 'convert_to_images':
                    # Convert to images
                    images = convert_from_path(file_path, dpi=200)
                    temp_dir = tempfile.mkdtemp()
                    
                    image_paths = []
                    for i, image in enumerate(images):
                        image_path = os.path.join(temp_dir, f"{pdf_file.name}_page_{i+1}.png")
                        image.save(image_path, 'PNG')
                        image_paths.append(image_path)
                    
                    results.append({
                        'filename': pdf_file.name,
                        'operation': 'convert_to_images',
                        'result': 'success',
                        'total_images': len(images),
                        'image_paths': image_paths
                    })
                
                # Add more operations as needed
                
            except Exception as file_error:
                results.append({
                    'filename': pdf_file.name,
                    'operation': operation,
                    'result': 'error',
                    'error': str(file_error)
                })
        
        return JsonResponse({
            'success': True,
            'operation': operation,
            'total_files': len(pdf_files),
            'results': results
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@require_http_methods(["GET"])
def get_pdf_metadata(request):
    """Get detailed metadata from PDF"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        metadata = {}
        
        # Basic metadata using PyPDF2
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            
            if reader.metadata:
                metadata.update({
                    'title': reader.metadata.get('/Title', 'Unknown'),
                    'author': reader.metadata.get('/Author', 'Unknown'),
                    'subject': reader.metadata.get('/Subject', 'Unknown'),
                    'creator': reader.metadata.get('/Creator', 'Unknown'),
                    'producer': reader.metadata.get('/Producer', 'Unknown'),
                    'creation_date': str(reader.metadata.get('/CreationDate', 'Unknown')),
                    'modification_date': str(reader.metadata.get('/ModDate', 'Unknown')),
                })
            
            metadata.update({
                'pages': len(reader.pages),
                'encrypted': reader.is_encrypted,
                'file_size': os.path.getsize(file_path)
            })
        
        # Advanced metadata using PyMuPDF
        try:
            doc = fitz.open(file_path)
            metadata.update({
                'format': 'PDF',
                'version': doc.pdf_version(),
                'page_count': doc.page_count,
                'needs_pass': doc.needs_pass,
                'permissions': doc.permissions,
                'language': doc.metadata.get('language', 'Unknown'),
                'keywords': doc.metadata.get('keywords', 'Unknown'),
            })
            
            # Get page dimensions
            page_sizes = []
            for page_num in range(min(5, doc.page_count)):  # First 5 pages
                page = doc.load_page(page_num)
                rect = page.rect
                page_sizes.append({
                    'page': page_num + 1,
                    'width': round(rect.width, 2),
                    'height': round(rect.height, 2),
                    'rotation': page.rotation
                })
            
            metadata['page_sizes'] = page_sizes
            doc.close()
            
        except Exception as advanced_error:
            metadata['advanced_error'] = str(advanced_error)
        
        return JsonResponse({
            'success': True,
            'metadata': metadata
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def optimize_pdf(request):
    """Optimize PDF for web or print"""
    try:
        file_path = request.session.get('current_pdf')
        if not file_path or not os.path.exists(file_path):
            return JsonResponse({'error': 'No PDF file found'}, status=400)
        
        data = json.loads(request.body)
        optimization_type = data.get('type', 'web')  # 'web' or 'print'
        
        doc = fitz.open(file_path)
        
        temp_dir = tempfile.mkdtemp()
        optimized_filename = f"optimized_{optimization_type}_{request.session.get('original_filename', 'file.pdf')}"
        optimized_path = os.path.join(temp_dir, optimized_filename)
        
        if optimization_type == 'web':
            # Optimize for web (smaller file size, lower quality)
            doc.save(
                optimized_path,
                garbage=4,
                deflate=True,
                clean=True,
                ascii=True,
                linear=True,
                pretty=False
            )
        else:
            # Optimize for print (better quality, larger file size)
            doc.save(
                optimized_path,
                garbage=3,
                deflate=True,
                clean=True,
                pretty=True
            )
        
        doc.close()
        
        # Get file sizes
        original_size = os.path.getsize(file_path)
        optimized_size = os.path.getsize(optimized_path)
        size_reduction = (original_size - optimized_size) / original_size * 100
        
        return JsonResponse({
            'success': True,
            'optimized_file': {
                'filename': optimized_filename,
                'path': optimized_path
            },
            'optimization_type': optimization_type,
            'original_size': original_size,
            'optimized_size': optimized_size,
            'size_reduction': round(size_reduction, 2)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# Cleanup function
def cleanup_temp_files(request):
    """Clean up temporary files (call this periodically or on session end)"""
    try:
        # This would typically be called by a cleanup task or session middleware
        temp_files = request.session.get('temp_files', [])
        
        for file_path in temp_files:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
        
        request.session['temp_files'] = []
        
        return JsonResponse({'success': True, 'cleaned_files': len(temp_files)})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)